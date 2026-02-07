import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DisputeWriter:
    def __init__(self, output_dir="pocket_credit/output_letters"):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def generate_round1(self, user_info, negative_items):
        """Round 1: Validation of Debt - Initial Dispute (Combined Letter)"""
        doc = Document()
        
        # Title
        title = doc.add_heading('DISPUTE LETTER - REQUEST FOR VALIDATION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_header(doc, user_info)
        
        doc.add_paragraph("To Whom It May Concern:")
        
        # Legal Introduction
        intro = doc.add_paragraph()
        intro.add_run(
            "I am writing to formally dispute the accuracy of the information listed below, "
            "which appears on my credit report. Pursuant to the Fair Credit Reporting Act "
            "(FCRA), 15 U.S.C. § 1681i, you are required to investigate and correct any inaccurate information."
        )
        
        # Legal Demand
        demand = doc.add_paragraph()
        demand.add_run(
            "Under FCRA Section 609 (15 U.S.C. § 1681g), I am entitled to request verification of any "
            "information that you have reported. Under FCRA Section 611 (15 U.S.C. § 1681i), you must "
            "conduct a reasonable investigation within 30 days of receiving this dispute."
        )
        
        # Items
        self._add_items_detailed(doc, negative_items)
        
        # Legal Warning
        warning = doc.add_paragraph()
        warning.add_run("LEGAL NOTICE: ").bold = True
        warning.add_run(
            "Failure to respond within 30 days constitutes a violation of the FCRA. "
            "Under 15 U.S.C. § 1681n and § 1681o, I may be entitled to actual damages, statutory damages "
            "of $100 to $1,000 per violation, punitive damages, and attorney's fees."
        )
        
        # Request
        doc.add_paragraph(
            "I request that you:\n"
            "1. Investigate the disputed items\n"
            "2. Provide me with copies of any documentation used to verify the accuracy\n"
            "3. Delete any information that cannot be verified\n"
            "4. Provide me with an updated copy of my credit report"
        )
        
        self._add_closing(doc, user_info)
        return self._save(doc, "Round1_Validation_")

    def batch_generate_round1(self, user_info, negative_items):
        """
        Generate INDIVIDUAL dispute letters - one per creditor/item.
        Returns list of file paths.
        """
        generated_letters = []
        
        # Get specific legal citations based on account type
        def get_legal_citations(item):
            creditor = item.get('creditor', '').upper()
            status = item.get('status', '').upper()
            
            citations = {
                'base': [
                    ('FCRA § 609', '15 U.S.C. § 1681g', 'Right to disclosure of information'),
                    ('FCRA § 611', '15 U.S.C. § 1681i', 'Duty to investigate disputed items'),
                    ('FCRA § 623', '15 U.S.C. § 1681s-2', 'Duty of furnishers to provide accurate information'),
                ]
            }
            
            # Add specific citations based on account type
            if 'COLLECTION' in status or 'COLLECTION' in creditor:
                citations['specific'] = [
                    ('FDCPA § 809', '15 U.S.C. § 1692g', 'Validation of debts - collector must provide verification'),
                    ('FDCPA § 807', '15 U.S.C. § 1692e', 'False or misleading representations prohibited'),
                ]
                citations['dispute_reason'] = (
                    "This collection account is disputed because:\n"
                    "• No original signed contract or agreement has been provided\n"
                    "• Chain of title/assignment documentation is missing\n"
                    "• The amount reported may include unauthorized fees or interest"
                )
            elif 'CHARGE' in status or 'LATE' in status:
                citations['specific'] = [
                    ('FCRA § 605', '15 U.S.C. § 1681c', 'Requirements relating to information in reports'),
                    ('FCRA § 607', '15 U.S.C. § 1681e', 'Compliance procedures - accuracy requirement'),
                ]
                citations['dispute_reason'] = (
                    "This negative payment history is disputed because:\n"
                    "• The late payment may have been reported in error\n"
                    "• Payment records may not have been properly credited\n"
                    "• The date and amount of alleged delinquency may be inaccurate"
                )
            elif 'REPO' in status:
                citations['specific'] = [
                    ('UCC § 9-614', '', 'Contents and form of notification before disposition'),
                    ('UCC § 9-616', '', 'Explanation of calculation of surplus or deficiency'),
                ]
                citations['dispute_reason'] = (
                    "This repossession account is disputed because:\n"
                    "• Proper notification may not have been provided before sale\n"
                    "• The vehicle may not have been sold in commercially reasonable manner\n"
                    "• The deficiency balance calculation may be incorrect"
                )
            else:
                citations['specific'] = []
                citations['dispute_reason'] = (
                    "This account is disputed because:\n"
                    "• The accuracy of the reported information cannot be verified\n"
                    "• There may be errors in the account details, dates, or amounts\n"
                    "• I do not have records confirming this account is accurate"
                )
            
            return citations
        
        for idx, item in enumerate(negative_items, 1):
            doc = Document()
            creditor = item.get('creditor', 'Unknown')[:30]
            account_num = item.get('account_num', 'Unknown')
            balance = item.get('balance', '0')
            citations = get_legal_citations(item)
            
            # Title
            title = doc.add_heading(f'DISPUTE LETTER - {creditor}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f"Account: {account_num} | Reported Balance: ${balance}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self._add_header(doc, user_info)
            
            # Address to bureau
            to_para = doc.add_paragraph()
            to_para.add_run("TO: Credit Bureau Compliance Department\n").bold = True
            to_para.add_run("RE: ").bold = True
            to_para.add_run(f"Formal Dispute of Account: {creditor}\n")
            to_para.add_run(f"Account Number: {account_num}\n")
            to_para.add_run(f"Reported Balance: ${balance}")
            
            doc.add_paragraph()
            doc.add_paragraph("To Whom It May Concern:")
            
            # Introduction
            intro = doc.add_paragraph()
            intro.add_run(
                f"I am writing to formally dispute the account listed above which appears on my "
                f"credit report. I am exercising my rights under federal law to demand verification "
                f"of this account and request immediate investigation."
            )
            
            # Legal Authority Section
            legal_heading = doc.add_heading('LEGAL AUTHORITY', level=1)
            
            legal_para = doc.add_paragraph()
            legal_para.add_run("This dispute is made pursuant to the following federal laws:\n\n").bold = True
            
            # Base citations
            for cite in citations['base']:
                legal_para.add_run(f"• {cite[0]} ({cite[1]}): ").bold = True
                legal_para.add_run(f"{cite[2]}\n")
            
            # Specific citations
            for cite in citations.get('specific', []):
                legal_para.add_run(f"• {cite[0]} ({cite[1]}): ").bold = True
                legal_para.add_run(f"{cite[2]}\n")
            
            # Dispute Reason
            reason_heading = doc.add_heading('REASON FOR DISPUTE', level=1)
            reason_para = doc.add_paragraph()
            reason_para.add_run(citations.get('dispute_reason', 'Information is inaccurate'))
            
            # Specific Demands
            demands_heading = doc.add_heading('DEMANDS', level=1)
            demands = doc.add_paragraph()
            demands.add_run(
                "I hereby demand that you:\n\n"
                f"1. Conduct a thorough investigation of this {creditor} account\n"
                "2. Contact the data furnisher and require them to verify:\n"
                "   • Original signed contract or agreement\n"
                "   • Complete payment history\n"
                "   • Accurate balance calculation\n"
                "   • Proof of ownership/assignment (if applicable)\n"
                "3. Provide me with all documentation used in your investigation\n"
                "4. DELETE this account if it cannot be verified with documentation\n"
                "5. Send me an updated credit report showing the results"
            )
            
            # Legal Warning
            warning_heading = doc.add_heading('LEGAL NOTICE', level=1)
            warning_para = doc.add_paragraph()
            warning_para.add_run("YOU HAVE 30 DAYS TO RESPOND\n\n").bold = True
            warning_para.add_run(
                "Under FCRA § 611, you must complete your investigation within 30 days. "
                "Failure to do so requires immediate deletion of this item.\n\n"
                "Under 15 U.S.C. § 1681n, willful non-compliance may result in:\n"
                "• Statutory damages of $100 - $1,000\n"
                "• Punitive damages\n"
                "• Attorney's fees and costs\n\n"
                "I am keeping detailed records of all correspondence. "
                "If this dispute is not resolved satisfactorily, I am prepared to escalate "
                "to the CFPB, FTC, State Attorney General, and/or federal court."
            )
            
            self._add_closing(doc, user_info)
            
            # Save with creditor name in filename
            safe_name = ''.join(c for c in creditor if c.isalnum() or c == ' ')[:20].replace(' ', '_')
            filename = f"R1_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            path = os.path.join(self.output_dir, filename)
            doc.save(path)
            generated_letters.append(path)
            print(f"  ✅ Generated: {filename}")
        
        return generated_letters

    def generate_round2(self, user_info, items):
        """Round 2: Method of Verification (609 Request)"""
        doc = Document()
        
        title = doc.add_heading('DEMAND FOR METHOD OF VERIFICATION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_header(doc, user_info)
        
        doc.add_paragraph("To Compliance Department:")
        
        # Reference previous dispute
        ref = doc.add_paragraph()
        ref.add_run("RE: CONTINUED DISPUTE - SECOND NOTICE\n").bold = True
        ref.add_run(
            "I previously disputed the items listed below. Your response indicated these items were "
            "'verified as accurate.' However, you failed to provide any documentation or explain the "
            "method of verification as required by law."
        )
        
        # Legal Demand
        legal = doc.add_paragraph()
        legal.add_run("LEGAL REQUIREMENTS:\n").bold = True
        legal.add_run(
            "Under FCRA Section 611(a)(6)(B)(iii) [15 U.S.C. § 1681i(a)(6)(B)(iii)], you are required "
            "to provide me with a description of the procedure used to determine the accuracy of the "
            "disputed information. Simply stating 'verified' without explanation does not satisfy "
            "your legal obligations."
        )
        
        # Specific Questions
        questions = doc.add_paragraph()
        questions.add_run("I DEMAND THE FOLLOWING INFORMATION:\n").bold = True
        questions.add_run(
            "1. The name of the person who verified the disputed information\n"
            "2. The business address of the source furnisher\n"
            "3. The telephone number of the source furnisher\n"
            "4. The specific documents reviewed during verification\n"
            "5. The date the verification was completed"
        )
        
        self._add_items_detailed(doc, items)
        
        # Consequence Warning
        consequence = doc.add_paragraph()
        consequence.add_run("NOTICE OF CONSEQUENCES: ").bold = True
        consequence.add_run(
            "If you cannot provide this information, you are in violation of FCRA § 611 and these items "
            "must be immediately deleted from my credit file. Continued reporting of unverified information "
            "may constitute willful non-compliance under 15 U.S.C. § 1681n."
        )
        
        self._add_closing(doc, user_info)
        return self._save(doc, "Round2_MethodVerification_")

    def generate_round3(self, user_info, items):
        """Round 3: Warning of Non-Compliance"""
        doc = Document()
        
        title = doc.add_heading('FORMAL NOTICE OF FCRA VIOLATION', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self._add_header(doc, user_info)
        
        doc.add_paragraph("CERTIFIED MAIL - RETURN RECEIPT REQUESTED\n\nTo Legal/Compliance Department:")
        
        # Strong Opening
        opening = doc.add_paragraph()
        opening.add_run("RE: FORMAL COMPLAINT - WILLFUL NON-COMPLIANCE WITH FCRA\n").bold = True
        opening.add_run(
            "This letter serves as formal notice that you are in VIOLATION of federal law. "
            "Despite multiple disputes and requests for verification, you have failed to comply "
            "with your obligations under the Fair Credit Reporting Act."
        )
        
        # Violation Summary
        violation = doc.add_paragraph()
        violation.add_run("DOCUMENTED VIOLATIONS:\n").bold = True
        violation.add_run(
            "• Failure to conduct reasonable investigation (15 U.S.C. § 1681i(a)(1))\n"
            "• Failure to provide method of verification (15 U.S.C. § 1681i(a)(6)(B)(iii))\n"
            "• Continued reporting of inaccurate/unverifiable information (15 U.S.C. § 1681e(b))\n"
            "• Potential willful non-compliance (15 U.S.C. § 1681n)"
        )
        
        self._add_items_detailed(doc, items)
        
        # Regulatory Threats
        threats = doc.add_paragraph()
        threats.add_run("NOTICE OF INTENDED ACTIONS:\n").bold = True
        threats.add_run(
            "If these items are not DELETED within 15 days of receipt of this letter, I will:\n\n"
            "1. File a formal complaint with the Consumer Financial Protection Bureau (CFPB)\n"
            "2. File a complaint with the Federal Trade Commission (FTC)\n"
            "3. File a complaint with my State Attorney General's Office\n"
            "4. Consult with an FCRA attorney regarding litigation options\n"
            "5. Seek statutory damages of $1,000 per violation under 15 U.S.C. § 1681n"
        )
        
        # Deadline
        deadline = doc.add_paragraph()
        deadline.add_run(f"DEADLINE: {(datetime.now().replace(day=datetime.now().day) + __import__('datetime').timedelta(days=15)).strftime('%B %d, %Y')}").bold = True
        
        self._add_closing(doc, user_info)
        return self._save(doc, "Round3_ViolationNotice_")

    def generate_round4(self, user_info, items):
        """Round 4: Intent to Litigate"""
        doc = Document()
        
        title = doc.add_heading('NOTICE OF INTENT TO FILE LAWSUIT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph("FINAL NOTICE BEFORE LITIGATION")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].bold = True
        
        self._add_header(doc, user_info)
        
        doc.add_paragraph("CERTIFIED MAIL - RETURN RECEIPT REQUESTED\n\nTo: General Counsel / Legal Department:")
        
        # Final Warning
        warning = doc.add_paragraph()
        warning.add_run("THIS IS YOUR FINAL NOTICE.\n\n").bold = True
        warning.add_run(
            "You have ignored my repeated disputes and requests for verification regarding the accounts "
            "listed below. Your continued reporting of this inaccurate and unverifiable information is "
            "causing me ongoing financial harm, including denial of credit, higher interest rates, "
            "and damage to my reputation."
        )
        
        # Legal Basis
        legal = doc.add_paragraph()
        legal.add_run("CAUSES OF ACTION:\n").bold = True
        legal.add_run(
            "Based on your conduct, I intend to pursue the following claims:\n\n"
            "1. WILLFUL NON-COMPLIANCE (15 U.S.C. § 1681n)\n"
            "   - Statutory damages: $100 - $1,000 per violation\n"
            "   - Punitive damages as the court may allow\n"
            "   - Attorney's fees and costs\n\n"
            "2. NEGLIGENT NON-COMPLIANCE (15 U.S.C. § 1681o)\n"
            "   - Actual damages suffered\n"
            "   - Attorney's fees and costs\n\n"
            "3. DEFAMATION (State Law)\n"
            "   - Publication of false information to third parties"
        )
        
        self._add_items_detailed(doc, items)
        
        # Demand
        demand = doc.add_paragraph()
        demand.add_run("DEMAND:\n").bold = True
        demand.add_run(
            "You have FIVE (5) BUSINESS DAYS from receipt of this letter to:\n\n"
            "1. DELETE all disputed items from my credit file, and\n"
            "2. Provide written confirmation of deletion\n\n"
            "If I do not receive confirmation of deletion within this timeframe, I will proceed with "
            "filing a lawsuit in federal court without further notice."
        )
        
        # Preservation Notice
        preserve = doc.add_paragraph()
        preserve.add_run("LITIGATION HOLD NOTICE: ").bold = True
        preserve.add_run(
            "You are hereby directed to preserve all documents, records, communications, and materials "
            "related to my credit file and the disputed accounts. Failure to preserve evidence may result "
            "in adverse inference at trial."
        )
        
        self._add_closing(doc, user_info, include_attorney_note=True)
        return self._save(doc, "Round4_IntentToLitigate_")

    def _add_header(self, doc, user_info):
        """Add professional header with user info"""
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        
        from_info = doc.add_paragraph()
        from_info.add_run("FROM:\n").bold = True
        from_info.add_run(f"{user_info.get('name', 'Consumer')}\n")
        from_info.add_run(f"{user_info.get('address', '[Address]')}\n")
        from_info.add_run(f"SSN (Last 4): xxx-xx-{user_info.get('ssn', 'XXXX')}\n")
        from_info.add_run(f"DOB: {user_info.get('dob', '[DOB]')}")
        
        doc.add_paragraph("—" * 50)

    def _add_items_detailed(self, doc, items):
        """Add disputed items with full details"""
        heading = doc.add_heading('DISPUTED ACCOUNTS:', level=1)
        
        for idx, item in enumerate(items, 1):
            p = doc.add_paragraph()
            p.add_run(f"Account #{idx}\n").bold = True
            p.add_run(f"Creditor/Furnisher: {item.get('creditor', 'Unknown')}\n")
            p.add_run(f"Account Number: {item.get('account_num', 'Unknown')}\n")
            p.add_run(f"Reported Balance: ${item.get('balance', 'Unknown')}\n")
            p.add_run(f"Dispute Reason: Information is inaccurate, unverified, and/or obsolete\n")
            doc.add_paragraph()

    def _add_closing(self, doc, user_info, include_attorney_note=False):
        """Add professional closing"""
        doc.add_paragraph()
        
        closing = doc.add_paragraph()
        closing.add_run("Respectfully,\n\n\n")
        closing.add_run(f"_______________________________\n")
        closing.add_run(f"{user_info.get('name', 'Consumer')}")
        
        if include_attorney_note:
            note = doc.add_paragraph()
            note.add_run("\nCC: File (for litigation records)\n").italic = True
            note.add_run("This letter may be used as an exhibit in federal court.").italic = True

    def _save(self, doc, prefix):
        """Save document with timestamp"""
        filename = f"{prefix}{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = os.path.join(self.output_dir, filename)
        doc.save(path)
        return path


    def generate_ai_letter(self, user_info, item_data, custom_instructions=""):
        """
        Generate a highly customized dispute letter using Google Gemini.
        """
        doc = Document()
        
        # 1. Setup Gemini
        import os
        api_key = os.environ.get("GOOGLE_API_KEY")
        if not api_key:
             try:
                 import streamlit as st
                 if "GOOGLE_API_KEY" in st.secrets:
                     api_key = st.secrets["GOOGLE_API_KEY"]
             except:
                 pass
        
        if not api_key:
            return "ERROR: No Google API Key found. Go to Settings."

        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            # 2. Construct Prompt
            creditor = item_data.get('creditor', 'Unknown')
            prompt = f"""
            Write a formal, legally aggressive credit dispute letter.
            
            **Sender:** {user_info.get('name')}
            **Recipient:** Credit Bureau Compliance Dept.
            **Account Disputed:** {creditor} (Account #: {item_data.get('account_num')})
            **Balance:** {item_data.get('balance')}
            **Context:** {custom_instructions}
            
            **Requirements:**
            - Use 15 U.S.C. 1681 citations (FCRA).
            - Be specific about why this is being disputed (if context provided).
            - Demand investigation or deletion.
            - Professional, firm legal tone.
            - NO placeholders. Fill everything out.
            - Format: Plain text, standard business letter format.
            """
            
            response = model.generate_content(prompt)
            letter_text = response.text
            
            # 3. Add content to Docx
            # Clean up asterisks often used by AI
            clean_text = letter_text.replace('**', '').replace('*', '')
            
            title = doc.add_heading('CUSTOM DISPUTE - AI DRAFTED', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self._add_header(doc, user_info)
            
            # Add AI content as paragraphs
            for para in clean_text.split('\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            return self._save(doc, f"AI_Dispute_{creditor.replace(' ','')}_")

        except Exception as e:
            return f"Error using Gemini: {str(e)}"

if __name__ == "__main__":
    # Test
    writer = DisputeWriter()
    test_user = {"name": "Test User", "address": "123 Test St", "ssn": "1234", "dob": "01/01/1990"}
    test_items = [{"creditor": "Test Creditor", "account_num": "12345", "balance": "500"}]
    print(writer.generate_round1(test_user, test_items))
