import streamlit as st
import sys
import os
import json
from datetime import datetime, timedelta
import shutil
import subprocess
from docx import Document

def get_docx_text(path):
    """Extract text from DOCX file for preview"""
    try:
        doc = Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading document: {e}"

# Add tools to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'tools'))
from parser import parse_credit_report
from writer import DisputeWriter
import client_manager as cm

# Configuration
BASE_DIR = os.path.dirname(__file__)
INPUT_DIR = os.path.join(BASE_DIR, "input_reports")
OUTPUT_DIR = os.path.join(BASE_DIR, "output_letters")
DISPUTES_FILE = os.path.join(BASE_DIR, "disputes.json")
PERSONAL_FILE = os.path.join(BASE_DIR, "personal_credit.json")


os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
ARCHIVE_DIR = os.path.join(BASE_DIR, "archive_letters")
os.makedirs(ARCHIVE_DIR, exist_ok=True)

# Note: st.set_page_config moved to __main__ block

# Increase file upload size limit (default is 200MB, we allow 500MB)
# Note: This requires setting via streamlit config or command line

# --- Personal Credit Data ---
def load_personal():
    if os.path.exists(PERSONAL_FILE):
        with open(PERSONAL_FILE, 'r') as f:
            return json.load(f)
    return {
        "name": "CeeJay",
        "address": "",
        "ssn": "0000",
        "dob": "",
        "disputes": [],
        "negative_items": []
    }

def save_personal(data):
    with open(PERSONAL_FILE, 'w') as f:
        json.dump(data, f, indent=2)

# --- Dispute Tracker Logic ---
def load_disputes():
    if os.path.exists(DISPUTES_FILE):
        with open(DISPUTES_FILE, 'r') as f:
            return json.load(f)
    return []

def save_dispute(client_name, letter_path, items, is_personal=False):
    if is_personal:
        personal = load_personal()
        personal["disputes"].append({
            "date_sent": datetime.now().strftime("%Y-%m-%d"),
            "letter_path": letter_path,
            "items": [i.get('creditor', 'Unknown') for i in items],
            "status": "Sent",
            "next_action_date": (datetime.now() + timedelta(days=35)).strftime("%Y-%m-%d")
        })
        save_personal(personal)
    else:
        disputes = load_disputes()
        disputes.append({
            "client": client_name,
            "date_sent": datetime.now().strftime("%Y-%m-%d"),
            "letter_path": letter_path,
            "items": [i.get('creditor', 'Unknown') for i in items],
            "status": "Sent",
            "next_action_date": (datetime.now() + timedelta(days=35)).strftime("%Y-%m-%d")
        })
        with open(DISPUTES_FILE, 'w') as f:
            json.dump(disputes, f, indent=2)

def main():
    # --- Sidebar: Mode Selector ---
    st.sidebar.title("💳 Credit Repair")
    
    mode = st.sidebar.radio("Mode", ["🔒 My Personal Credit", "👥 Client Mode"], index=0)
    
    personal = load_personal()
    current_client = None
    
    if mode == "🔒 My Personal Credit":
        st.sidebar.success(f"Personal Mode: {personal['name']}")
        
        # Personal Settings
        with st.sidebar.expander("⚙️ Edit My Info"):
            new_name = st.text_input("Name", value=personal["name"])
            new_addr = st.text_area("Address", value=personal["address"])
            new_ssn = st.text_input("SSN (Last 4)", value=personal["ssn"], type="password")
            new_dob = st.text_input("DOB", value=personal["dob"])
            if st.button("Save"):
                personal["name"] = new_name
                personal["address"] = new_addr
                personal["ssn"] = new_ssn
                personal["dob"] = new_dob
                save_personal(personal)
                st.success("Saved!")
                st.rerun()
        
        current_client = personal
        is_personal = True
    else:
        is_personal = False
        st.sidebar.subheader("👥 Client Manager")
        
        clients = cm.load_clients()
        client_names = [c["name"] for c in clients]
        selected_client_name = st.sidebar.selectbox("Select Client", ["New Client..."] + client_names)
    
        if selected_client_name == "New Client...":
            with st.sidebar.form("new_client"):
                name = st.text_input("Full Name")
                addr = st.text_area("Address")
                ssn = st.text_input("SSN (Last 4)", type="password")
                dob = st.text_input("DOB (MM/DD/YYYY)")
                if st.form_submit_button("Save Client"):
                    if name:
                        cm.add_client(name, addr, ssn, dob)
                        st.success("Client Saved!")
                        st.rerun()
        else:
            current_client = cm.get_client(selected_client_name)
            st.sidebar.success(f"Active: {current_client['name']}")
            
            if st.sidebar.button("🗑️ Delete Client"):
                cm.delete_client(selected_client_name)
                st.rerun()
    
    # --- Main Content ---
    st.title("💳 Credit Repair Specialist")
    if is_personal:
        st.caption(f"📍 Personal Mode — Managing {personal['name']}'s Credit")
    else:
        st.caption("👥 Client Mode — Multi-Client Dispute Management")
    
    if not current_client:
        st.warning("Please select or create a client in the sidebar to begin.")
        st.stop()
    
    # Tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Dispute Tracker", "📤 Upload Reports", "📋 Negative Items", "✉️ Generated Letters", "⚙️ Settings"])
    
    # ========== TAB 1: Dispute Tracker with Countdown ==========
    with tab1:
        st.subheader(f"🔥 Active Disputes: {current_client['name']}")
        
        if is_personal:
            disputes = personal.get("disputes", [])
        else:
            disputes = [d for d in load_disputes() if d.get('client') == current_client['name']]
        
        if disputes:
            disputes.sort(key=lambda x: x['date_sent'], reverse=True)
            
            # Summary Stats
            col1, col2, col3 = st.columns(3)
            col1.metric("📨 Total Disputes Sent", len(disputes))
            
            pending = sum(1 for d in disputes if d['status'] == 'Sent')
            col2.metric("⏳ Pending Response", pending)
            
            urgent = sum(1 for d in disputes if (datetime.strptime(d['next_action_date'], "%Y-%m-%d") - datetime.now()).days <= 5)
            col3.metric("🔴 Urgent (< 5 days)", urgent)
            
            st.divider()
            
            for idx, d in enumerate(disputes):
                with st.container(border=True):
                    target_date = datetime.strptime(d['next_action_date'], "%Y-%m-%d")
                    days_left = (target_date - datetime.now()).days
                    
                    c1, c2, c3, c4 = st.columns([2, 1, 1, 1])
                    
                    c1.markdown(f"**{d['date_sent']}** — {len(d['items'])} Items")
                    c1.caption(f"📄 {os.path.basename(d['letter_path'])}")
                    
                    # Countdown Timer
                    if days_left > 10:
                        c2.success(f"✅ **{days_left} Days** left")
                    elif days_left > 5:
                        c2.warning(f"⚠️ **{days_left} Days** left")
                    elif days_left > 0:
                        c2.error(f"🔴 **{days_left} Days** left")
                    else:
                        c2.error(f"💥 **OVERDUE** by {abs(days_left)} days!")
                    
                    # Status Selector
                    status_options = ["Sent", "In Progress", "Responded - Need Review", "Deleted!", "No Response"]
                    current_status = d.get('status', 'Sent')
                    new_status = c3.selectbox("Status", status_options, index=status_options.index(current_status) if current_status in status_options else 0, key=f"status_{d['date_sent']}_{d['letter_path']}_{idx}")
                    
                    # Update Status Button
                    if new_status != current_status:
                        if c4.button("Update", key=f"upd_{d['date_sent']}_{idx}"):
                            d['status'] = new_status
                            if is_personal:
                                save_personal(personal)
                            else:
                                with open(DISPUTES_FILE, 'w') as f:
                                    json.dump(load_disputes(), f, indent=2)
                            st.rerun()
                    
                    # Delete Dispute Button
                    if c3.button("🗑️", key=f"del_disp_{d['date_sent']}_{idx}"):
                        if is_personal:
                             personal['disputes'].remove(d)
                             save_personal(personal)
                        else:
                             disputes.remove(d) # Need to reload and save strictly
                             # Re-load to be safe
                             all_disputes = load_disputes()
                             # Filter out this exact one
                             new_disputes = [x for x in all_disputes if not (x['client'] == d['client'] and x['date_sent'] == d['date_sent'] and x['letter_path'] == d['letter_path'])]
                             with open(DISPUTES_FILE, 'w') as f:
                                 json.dump(new_disputes, f, indent=2)
                        st.success("Deleted!")
                        st.rerun()
                    
                    # Show disputed items
                    with st.expander("View Disputed Items"):
                        for item in d['items']:
                            st.write(f"• {item}")
        else:
            st.info("No active disputes yet. Upload a credit report to get started!")
    
    # ========== TAB 2: Upload Reports (All 3 Bureaus) ==========
    with tab2:
        st.subheader("📤 Upload Credit Reports")
        st.caption("Upload reports from all 3 bureaus to dispute with each one separately.")
        
        # Initialize bureau data in personal/session
        if is_personal:
            if 'bureaus' not in personal:
                personal['bureaus'] = {'experian': [], 'equifax': [], 'transunion': []}
                save_personal(personal)
            bureaus_data = personal['bureaus']
        else:
            if 'bureaus' not in st.session_state:
                st.session_state['bureaus'] = {'experian': [], 'equifax': [], 'transunion': []}
            bureaus_data = st.session_state['bureaus']
        
        # Summary Cards
        col1, col2, col3 = st.columns(3)
        col1.metric("🔵 Experian", f"{len(bureaus_data.get('experian', []))} items")
        col2.metric("🔴 Equifax", f"{len(bureaus_data.get('equifax', []))} items")
        col3.metric("🟢 TransUnion", f"{len(bureaus_data.get('transunion', []))} items")
        
        st.divider()
        
        # Bureau Upload Tabs + Camera Scanner
        exp_tab, eqf_tab, tu_tab, scan_tab = st.tabs(["🔵 Experian", "🔴 Equifax", "🟢 TransUnion", "📸 Scan with Camera"])
        
        def process_bureau_upload(bureau_name, bureau_key, uploaded_file):
            """Process upload for a specific bureau"""
            if uploaded_file:
                # Track processed files to avoid reprocessing on every rerun
                if 'processed_files' not in st.session_state:
                    st.session_state['processed_files'] = set()
                
                file_id = f"{bureau_key}_{uploaded_file.name}_{uploaded_file.size}"
                
                # Skip if already processed
                if file_id in st.session_state['processed_files']:
                    return
                
                save_path = os.path.join(INPUT_DIR, f"{bureau_key}_{uploaded_file.name}")
                with open(save_path, 'wb') as f:
                    f.write(uploaded_file.getbuffer())
                st.success(f"✅ Uploaded {bureau_name} report: {uploaded_file.name}")
                
                with st.spinner(f"Analyzing {bureau_name} report..."):
                    try:
                        items = parse_credit_report(save_path)
                        
                        # Mark as processed BEFORE saving to prevent loop
                        st.session_state['processed_files'].add(file_id)
                        
                        if items:
                            # Tag each item with its bureau
                            for item in items:
                                item['bureau'] = bureau_key
                            
                            st.success(f"Found **{len(items)}** negative items on {bureau_name}!")
                            
                            # Save to appropriate storage
                            if is_personal:
                                personal['bureaus'][bureau_key] = items
                                # Also update combined list
                                all_items = []
                                for b in ['experian', 'equifax', 'transunion']:
                                    all_items.extend(personal['bureaus'].get(b, []))
                                personal['negative_items'] = all_items
                                save_personal(personal)
                            else:
                                st.session_state['bureaus'][bureau_key] = items
                                all_items = []
                                for b in ['experian', 'equifax', 'transunion']:
                                    all_items.extend(st.session_state['bureaus'].get(b, []))
                                st.session_state['negative_items'] = all_items
                            
                            # Show preview
                            with st.expander(f"View {bureau_name} Items"):
                                for item in items:
                                    st.write(f"• **{item.get('creditor', 'Unknown')}** - ${item.get('balance', '0')}")
                            
                            st.rerun()
                        else:
                            st.success(f"🎉 No negative items on {bureau_name}! It's clean!")
                    except Exception as e:
                        st.error(f"Error parsing {bureau_name} report: {e}")
        
        with exp_tab:
            st.markdown("### 🔵 Experian Report")
            exp_items = bureaus_data.get('experian', [])
            if exp_items:
                st.success(f"✅ {len(exp_items)} items loaded from Experian")
                if st.button("🗑️ Clear Experian Data", key="clear_exp"):
                    if is_personal:
                        personal['bureaus']['experian'] = []
                        save_personal(personal)
                    else:
                        st.session_state['bureaus']['experian'] = []
                    st.rerun()
            exp_upload = st.file_uploader("Upload Experian (PDF, Pages, DOCX, TXT, Image)", type=['pdf', 'pages', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg'], key="exp_upload")
            process_bureau_upload("Experian", "experian", exp_upload)
        
        with eqf_tab:
            st.markdown("### 🔴 Equifax Report")
            eqf_items = bureaus_data.get('equifax', [])
            if eqf_items:
                st.success(f"✅ {len(eqf_items)} items loaded from Equifax")
                if st.button("🗑️ Clear Equifax Data", key="clear_eqf"):
                    if is_personal:
                        personal['bureaus']['equifax'] = []
                        save_personal(personal)
                    else:
                        st.session_state['bureaus']['equifax'] = []
                    st.rerun()
            eqf_upload = st.file_uploader("Upload Equifax (PDF, Pages, DOCX, TXT, Image)", type=['pdf', 'pages', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg'], key="eqf_upload")
            process_bureau_upload("Equifax", "equifax", eqf_upload)
        
        with tu_tab:
            st.markdown("### 🟢 TransUnion Report")
            tu_items = bureaus_data.get('transunion', [])
            if tu_items:
                st.success(f"✅ {len(tu_items)} items loaded from TransUnion")
                if st.button("🗑️ Clear TransUnion Data", key="clear_tu"):
                    if is_personal:
                        personal['bureaus']['transunion'] = []
                        save_personal(personal)
                    else:
                        st.session_state['bureaus']['transunion'] = []
                    st.rerun()
            tu_upload = st.file_uploader("Upload TransUnion (PDF, Pages, DOCX, TXT, Image)", type=['pdf', 'pages', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg'], key="tu_upload")
            process_bureau_upload("TransUnion", "transunion", tu_upload)
        


# ... (Inside tab2)
        with scan_tab:

            st.markdown("### 📸 Scan Credit Report with Camera")
            st.caption("Take a photo of your credit report page. AI will extract the negative accounts.")
            
            bureau_choice = st.selectbox("Which bureau is this from?", ["Experian", "Equifax", "TransUnion"], key="scan_bureau")
            bureau_key = bureau_choice.lower()
            
            # Input Method Toggle
            input_method = st.radio("Input Method", ["📁 Upload (Scanner/Photo)", "📸 Live Camera"], horizontal=True)
            
            camera_img = None
            if input_method == "📸 Live Camera":
                camera_img = st.camera_input("Take Photo", key="webcam_scan")
            else:
                camera_img = st.file_uploader("Upload Image", type=['png', 'jpg', 'jpeg', 'heic'], accept_multiple_files=False, key="mobile_scan")

            if camera_img:
                # ... (Existing analysis logic remains the same)
                st.image(camera_img, caption="Uploaded Image", use_container_width=True)
                
                if st.button("🧠 Analyze with AI", type="primary"):
                    with st.spinner("AI is reading your credit report..."):
                        try:
                            import google.generativeai as genai
                            import base64
                            
                            api_key = st.secrets.get("GEMINI_API_KEY", "")
                            if not api_key:
                                st.error("Please set your Gemini API key in Settings first.")
                            else:
                                genai.configure(api_key=api_key)
                                model = genai.GenerativeModel("gemini-2.5-flash")
                                
                                img_bytes = camera_img.read()
                                img_b64 = base64.b64encode(img_bytes).decode()
                                
                                prompt = """Analyze this credit report image. Extract ALL negative accounts.
For each account, provide:
- Creditor Name
- Account Number (last 4 if visible)
- Balance
- Status (Collection, Charge-Off, Late, etc.)
- Date Opened (if visible)

Return as a JSON array like:
[{"creditor": "...", "account_number": "...", "balance": "...", "status": "...", "date_opened": "..."}]

If no negative accounts found, return empty array: []"""
                                
                                response = model.generate_content([
                                    {"mime_type": "image/jpeg", "data": img_b64}, # Default to jpeg for gemini compatibility
                                    prompt
                                ])
                                
                                import json
                                import re
                                
                                text = response.text
                                json_match = re.search(r'\[.*\]', text, re.DOTALL)
                                if json_match:
                                    items = json.loads(json_match.group())
                                    if items:
                                        st.success(f"Found {len(items)} accounts!")
                                        for item in items:
                                            item['bureau'] = bureau_key
                                            item['source'] = 'camera_scan'
                                        
                                        # Save to bureaus data
                                        if is_personal:
                                            personal['bureaus'][bureau_key].extend(items)
                                            # Also add to negative items main list if desired, or let user do it
                                            # For now, just bureau bucket
                                            save_personal(personal)
                                        else:
                                            st.session_state['bureaus'][bureau_key].extend(items)
                                        
                                        st.json(items)
                                        st.rerun()
                                    else:
                                        st.info("No negative accounts found in this image.")
                                else:
                                    st.warning("Could not parse accounts. Try a clearer photo.")
                                    st.text(text)
                        except Exception as e:
                            st.error(f"Error: {e}")

    # ========== TAB 4: Generated Letters ==========
    with tab4:
        st.subheader("✉️ Your Letters & Templates")
        
        # Sub-tabs for organization
        ai_tab, letters_tab, templates_tab = st.tabs(["✨ AI Specialist", "📄 Generated Letters", "📚 Templates"])
        
        with ai_tab:
            st.info("AI Specialist coming soon!")

        with letters_tab:
            letters = [f for f in os.listdir(OUTPUT_DIR) if f.endswith('.docx')]
            
            if letters:
                col_header1, col_header2 = st.columns([3, 1])
                col_header1.success(f"You have {len(letters)} letters generated!")
                
                # Clear All with Confirmation
                if "confirm_clear_all" not in st.session_state:
                    st.session_state.confirm_clear_all = False
                
                if st.session_state.confirm_clear_all:
                    col_header2.warning("Are you sure?")
                    c_yes, c_no = col_header2.columns(2)
                    if c_yes.button("Yes", key="yes_clear"):
                        for l in letters:
                             os.remove(os.path.join(OUTPUT_DIR, l))
                        st.session_state.confirm_clear_all = False
                        st.rerun()
                    if c_no.button("No", key="no_clear"):
                        st.session_state.confirm_clear_all = False
                        st.rerun()
                else:
                    if col_header2.button("🗑️ Clear All", type="secondary"):
                        st.session_state.confirm_clear_all = True
                        st.rerun()
                    
                for letter in sorted(letters, reverse=True):
                    letter_path = os.path.join(OUTPUT_DIR, letter)
                    
                    with st.container(border=True):
                        # Revised Columns: Name | Download | Print | Archive | Delete
                        c1, c2, c3, c4, c5 = st.columns([2, 0.6, 0.5, 0.5, 0.5])
                        c1.write(f"📄 **{letter}**")
                        
                        # Detect round from filename
                        if "Round1" in letter or letter.startswith("R1_"):
                            c1.caption("🔵 Round 1: Initial Validation")
                        elif "Round2" in letter:
                            c1.caption("🟡 Round 2: Method of Verification")
                        elif "Round3" in letter:
                            c1.caption("🟠 Round 3: Warning of Non-Compliance")
                        elif "Round4" in letter:
                            c1.caption("🔴 Round 4: Intent to Litigate")
                        
                        with open(letter_path, 'rb') as f:
                            c2.download_button("⬇️", f, file_name=letter, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_{letter}", help="Download")
                        
                        if c3.button("🖨️", key=f"prt_{letter}", help="Print"):
                            try:
                                rtf_path = letter_path.replace(".docx", ".rtf")
                                subprocess.run(["textutil", "-convert", "rtf", letter_path, "-output", rtf_path], check=True)
                                subprocess.run(["lp", rtf_path], check=True)
                                st.toast(f"Sent to printer!", icon="✅")
                                if os.path.exists(rtf_path):
                                    os.remove(rtf_path)
                            except Exception as e:
                                st.error(f"Print failed: {e}")
                                
                        if c4.button("📂", key=f"arch_{letter}", help="Archive"):
                            shutil.move(letter_path, os.path.join(ARCHIVE_DIR, letter))
                            st.toast("Archived!", icon="📦")
                            st.rerun()
                        
                        if c5.button("🗑️", key=f"del_{letter}", help="Delete"):
                            os.remove(letter_path)
                            st.rerun()
                    
                        # Preview Section
                        with st.expander("👁️ Preview"):
                            try:
                                preview_text = get_docx_text(letter_path)
                                st.text_area(f"Content", preview_text, height=200, disabled=True, key=f"prev_{letter}")
                            except Exception as e:
                                st.error(f"Preview failed: {e}") 
                               
            else:
                st.info("No letters generated yet. Go to 'Negative Items' tab to generate your first letter!")
        
        with templates_tab:
            st.markdown("### 📚 The 4 Dispute Letter Rounds")
            st.caption("This is your escalation strategy. Start with Round 1 and escalate if needed.")
            
            # Round 1
            with st.container(border=True):
                st.markdown("#### 🔵 Round 1: Validation of Debt (Initial Dispute)")
                st.write("**Legal Basis:** FCRA Section 609, 611")
                st.write("**Purpose:** Request the bureau to verify the accuracy of the reported items.")
                st.write("**Timeline:** Send first, then wait 30-35 days for response.")
                st.info("💡 **Tip:** Send certified mail with return receipt for proof.")
            
            # Round 2
            with st.container(border=True):
                st.markdown("#### 🟡 Round 2: Method of Verification (609 Request)")
                st.write("**Legal Basis:** FCRA Section 611(a)(6)(B)(iii)")
                st.write("**Purpose:** Demand to know HOW they verified the debt. If they can't explain, it must be removed.")
                st.write("**When to Use:** After Round 1 if items remain and they claim 'verified as accurate'.")
                st.warning("⚠️ **Key Question:** 'What method did you use to verify this debt?'")
            
            # Round 3  
            with st.container(border=True):
                st.markdown("#### 🟠 Round 3: Warning of Non-Compliance")
                st.write("**Legal Basis:** FCRA Willful Non-Compliance (Section 616, 617)")
                st.write("**Purpose:** Put them on notice that continued reporting is a legal violation.")
                st.write("**When to Use:** After Rounds 1 & 2 if they still haven't removed the items.")
                st.error("🔴 **Escalation:** Threaten CFPB and FTC complaints.")
            
            # Round 4
            with st.container(border=True):
                st.markdown("#### 🔴 Round 4: Intent to Litigate")
                st.write("**Legal Basis:** FCRA Section 616/617 - Statutory Damages")
                st.write("**Purpose:** Final notice before filing a lawsuit for $1,000+ per violation.")
                st.write("**When to Use:** LAST RESORT after all other rounds failed.")
                st.error("⚖️ **Warning:** Only send if you're prepared to follow through with legal action!")
            
            st.divider()
            st.markdown("### 📝 How to Generate Letters")
            st.write("1. Go to the **'📋 Negative Items'** tab")
            st.write("2. Select the dispute round from the dropdown")
            st.write("3. Click **'🚀 Generate Letter'**")
            st.write("4. Download and print from this tab!")
    
    # ========== TAB 5: Settings ==========
    with tab5:
        st.subheader("⚙️ Credit Repair Settings")
        
        if is_personal:
            st.markdown("### Your Personal Information")
            st.write(f"**Name:** {personal['name']}")
            st.write(f"**Address:** {personal['address'] or 'Not set'}")
            st.write(f"**SSN (Last 4):** ***-**-{personal['ssn']}")
            st.write(f"**DOB:** {personal['dob'] or 'Not set'}")
            
            st.divider()
            
            if st.button("🗑️ Clear All Personal Disputes"):
                personal['disputes'] = []
                save_personal(personal)
                st.success("Cleared!")
                st.rerun()
        else:
            st.info("Client settings managed in sidebar.")
    
    # ========== TAB 3: Negative Items ==========
    with tab3:
        st.subheader("📋 Negative Accounts & Collections")
        
        # Load Items (Personal or Client)
        if is_personal:
            items = personal.get('negative_items', [])
        else:
            items = st.session_state.get('negative_items', [])
            
        if items:
            st.success(f"Found {len(items)} negative items across all bureaus.")
            
            # Selection for Dispute
            selected_indices = []
            for i, item in enumerate(items):
                # Unique key: bureau + creditor + account + balance
                key = f"{item.get('bureau')}_{item.get('creditor')}_{item.get('account_number')}_{i}"
                if st.checkbox(f"{item.get('bureau').upper()} - {item.get('creditor')} (${item.get('balance', '0')})", key=key):
                    selected_indices.append(i)
            
            selected_items = [items[i] for i in selected_indices]
            
            st.divider()
            
            # --- AI Writer Section ---
            c_ai, c_manual = st.tabs(["✨ AI Smart Dispute", "📝 Manual Templates"])
            
            with c_ai:
                st.markdown("#### 🧠 AI Specialist Writer")
                st.caption("Let Gemini write a custom, highly specific dispute letter for these items.")
                
                strategy = st.selectbox("Dispute Strategy", [
                    "Factual Dispute (Wrong Balance/Date)",
                    "Validation of Debt (Show me the contract)",
                    "Metro 2 Compliance (e-OSCAR errors)",
                    "Identity Theft / Fraud",
                    "Goodwill Adjustment (Late Payment Removal)"
                ])
                
                if st.button("✨ Write Custom Letter", disabled=len(selected_items)==0):
                    with st.spinner("Consulting AI Specialist..."):
                        try:
                            import google.generativeai as genai
                            api_key = st.secrets.get("GEMINI_API_KEY", "")
                            
                            if not api_key:
                                st.error("Please set GEMINI_API_KEY in Settings!")
                            else:
                                genai.configure(api_key=api_key)
                                model = genai.GenerativeModel("gemini-2.5-flash")
                                
                                user_name = personal['name'] if is_personal else current_client['name']
                                user_addr = personal['address'] if is_personal else current_client['address']
                                
                                prompt = f"""
                                Write a formal credit dispute letter for {user_name}.
                                Address: {user_addr}
                                Date: {datetime.now().strftime('%B %d, %Y')}
                                
                                Strategy: {strategy}
                                
                                Items to Dispute:
                                {json.dumps(selected_items, indent=2)}
                                
                                Instructions:
                                1. Use a professional, firm legal tone.
                                2. Cite relevant FCRA/FDCPA laws based on the strategy.
                                3. Clearly list each account and the reason for dispute.
                                4. Demand deletion or correction.
                                5. Do NOT include placeholders like [Your Name], use the provided real data.
                                """
                                
                                response = model.generate_content(prompt)
                                st.session_state['ai_letter_content'] = response.text
                                st.success("Letter Drafted!")
                                
                        except Exception as e:
                            st.error(f"AI Error: {e}")
                
                # Show Draft
                if 'ai_letter_content' in st.session_state:
                    letter_text = st.text_area("Edit Draft", value=st.session_state['ai_letter_content'], height=400)
                    
                    if st.button("💾 Save as DOCX"):
                        doc = Document()
                        for line in letter_text.split('\n'):
                            doc.add_paragraph(line)
                        
                        filename = f"AI_Dispute_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                        save_path = os.path.join(OUTPUT_DIR, filename)
                        doc.save(save_path)
                        
                        # Save to tracker
                        save_dispute(
                            personal['name'] if is_personal else current_client['name'],
                            save_path,
                            selected_items,
                            is_personal
                        )
                        st.success(f"Saved to 'Generated Letters' tab as {filename}!")
                        
            with c_manual:
                st.markdown("#### 📄 Standard Templates")
                # Existing logic for standard templates...
                round_choice = st.selectbox("Select Round", ["Round 1: Validation", "Round 2: Method of Verification", "Round 3: Warning", "Round 4: Intent to Litigate", "Pay for Delete", "Goodwill Letter"])
                
                if st.button("🚀 Generate Template Letter", disabled=len(selected_items)==0):
                    if is_personal:
                        client_data = personal
                    else:
                        client_data = current_client
                        
                    # Map round to template file (simplified logic)
                    template_map = {
                        "Round 1: Validation": "round1_validation.docx",
                        "Round 2: Method of Verification": "round2_mov.docx",
                        "Round 3: Warning": "round3_warning.docx",
                        "Round 4: Intent to Litigate": "round4_litigate.docx",
                        "Pay for Delete": "pay_for_delete.docx",
                        "Goodwill Letter": "goodwill.docx"
                    }
                    
                    template_name = template_map.get(round_choice, "round1_validation.docx")
                    # In a real app, we'd use the template engine. For now, we'll create a basic DOCX.
                    
                    doc = Document()
                    doc.add_heading(f"Dispute Letter - {round_choice}", 0)
                    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
                    doc.add_paragraph(f"To: Credit Bureaus") # In real app, generate 3 letters
                    doc.add_paragraph(f"From: {client_data['name']}")
                    doc.add_paragraph(f"Address: {client_data['address']}")
                    doc.add_paragraph(f"SSN: ***-**-{client_data['ssn'][-4:]}")
                    doc.add_paragraph(f"DOB: {client_data['dob']}")
                    doc.add_paragraph("\n To Whom It May Concern,")
                    doc.add_paragraph("I am writing to dispute the following items on my credit report...")
                    
                    for item in selected_items:
                        doc.add_paragraph(f"Creditor: {item.get('creditor')}")
                        doc.add_paragraph(f"Account #: {item.get('account_number')}")
                        doc.add_paragraph(f"Reason: I demand validation of this debt.")
                        doc.add_paragraph("------------------------------------------------")
                    
                    doc.add_paragraph("Sincerely,")
                    doc.add_paragraph(client_data['name'])
                    
                    filename = f"R1_{client_data['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                    save_path = os.path.join(OUTPUT_DIR, filename)
                    doc.save(save_path)
                    
                    save_dispute(client_data['name'], save_path, selected_items, is_personal)
                    st.success(f"Generated {filename}! Go to 'Generated Letters' tab to download.")
                    
        else:
            st.info("No negative items found yet. Go to 'Upload Reports' tab first.")
