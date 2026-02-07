import json
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DATA_FILE = os.path.join(os.path.dirname(__file__), "invoices.json")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")

def load_data():
    """Load invoice data from JSON."""
    if not os.path.exists(DATA_FILE):
        local_data = {
            "invoices": [],
            "next_invoice_number": 1001,
            "company_info": {
                "name": "YOUR COMPANY NAME",
                "address": "Your Address", 
                "city_state_zip": "City, ST 00000",
                "phone": "(000) 000-0000",
                "email": "email@example.com"
            },
            "clients": []
        }
    else:
        with open(DATA_FILE, 'r') as f:
            local_data = json.load(f)

    # Try DB Overlay
    try:
        # Import here to avoid circular dependency issues if any
        import sys
        sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../pocket_core')))
        import db
        
        client = db.get_db()
        if client:
            # Fetch Invoices
            res = client.table("invoices").select("data").execute()
            db_invoices = [row['data'] for row in res.data if row.get('data')]
            
            # Merge: DB takes precedence or we just use DB if available?
            # For simplicity, if DB connects, we use DB invoices.
            # But we also need clients and company info.
            # Creating tables for those is best, but for now let's stick to local for settings/clients 
            # and DB for invoices (the main transactional data).
            
            # actually, let's just use local data structure but replace invoices list
            if db_invoices:
                local_data["invoices"] = db_invoices
                
    except Exception as e:
        # Fail silently to local
        pass
        
    return local_data

def save_data(data):
    """Save invoice data to JSON and Sync to DB."""
    # Local Save
    with open(DATA_FILE, 'w') as f:
        json.dump(data, f, indent=2, default=str)

    # DB Sync
    try:
        import sys
        # ensure path is there (it should be from load_data but good to be safe)
        sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../pocket_core')))
        import db
        
        client = db.get_db()
        if client:
            # Sync Invoices
            # We only sync items that are in the data object
            for inv in data.get("invoices", []):
                # Flatten what we can, put rest in data
                # We need to extract origin/dest/ref from nested 'details' if present, or it might be raw?
                # The invoice structure in create_invoice:
                # {invoice_number, client, line_items, total, ...}
                # It does NOT have 'details' key by default in my previous read of create_invoice
                # Wait, create_invoice in this file line 108:
                # invoice = {invoice_number, client, date, due_date, line_items...}
                # It does NOT have origin/dest/ref fields at top level! 
                # They must be in line items logic or notes? 
                
                # In dashboard.py (Step 5483), user inputs route origin/dest but where does it go?
                # It goes into line item description! "Load from Chicago to Miami..."
                # It is NOT structured!
                
                # So my SQL schema with origin/dest columns is actually unused unless I parse it.
                # That's fine. We rely on 'data' col.
                
                payload = {
                    "id": str(inv.get("invoice_number")), # ID is text in SQL
                    "client_name": inv.get("client", {}).get("name", "Unknown"),
                    "amount": inv.get("total", 0),
                    "status": inv.get("status", "draft"),
                    "data": inv,
                    "created_at": inv.get("created_at", datetime.now().isoformat())
                }
                client.table("invoices").upsert(payload).execute()
                
    except Exception as e:
        print(f"DB Sync Error: {e}")

def update_company_info(name, address, city_state_zip, phone, email):
    """Update company information."""
    data = load_data()
    data["company_info"] = {
        "name": name,
        "address": address,
        "city_state_zip": city_state_zip,
        "phone": phone,
        "email": email
    }
    save_data(data)
    return data["company_info"]

def add_client(name, address, city_state_zip, email=None, phone=None):
    """Add a new client to the client list."""
    data = load_data()
    
    client = {
        "id": len(data["clients"]) + 1,
        "name": name,
        "address": address,
        "city_state_zip": city_state_zip,
        "email": email,
        "phone": phone
    }
    
    data["clients"].append(client)
    save_data(data)
    return client

def get_clients():
    """Get all clients."""
    return load_data().get("clients", [])

def create_invoice(client_id, line_items, notes=None, due_days=30):
    """
    Create a new invoice.
    
    Args:
        client_id: ID of the client
        line_items: List of dicts with 'description', 'quantity', 'rate'
        notes: Optional notes to include
        due_days: Days until due (default 30)
    
    Returns:
        Invoice dict
    """
    data = load_data()
    
    # Find client
    client = None
    for c in data["clients"]:
        if c["id"] == client_id:
            client = c
            break
    
    if not client:
        raise ValueError(f"Client ID {client_id} not found")
    
    # Calculate totals
    subtotal = sum(item["quantity"] * item["rate"] for item in line_items)
    
    invoice = {
        "invoice_number": data["next_invoice_number"],
        "client": client,
        "date": datetime.now().isoformat(),
        "due_date": (datetime.now() + timedelta(days=due_days)).isoformat(),
        "line_items": line_items,
        "subtotal": subtotal,
        "total": subtotal,  # Add tax calculation here if needed
        "notes": notes,
        "status": "unpaid",
        "created_at": datetime.now().isoformat()
    }
    
    data["invoices"].append(invoice)
    data["next_invoice_number"] += 1
    save_data(data)
    
    return invoice

def get_invoices(status=None):
    """Get all invoices, optionally filtered by status."""
    data = load_data()
    invoices = data.get("invoices", [])
    
    if status:
        invoices = [i for i in invoices if i.get("status") == status]
    
    return sorted(invoices, key=lambda x: x["invoice_number"], reverse=True)

def mark_invoice_paid(invoice_number):
    """Mark an invoice as paid."""
    data = load_data()
    
    for inv in data["invoices"]:
        if inv["invoice_number"] == invoice_number:
            inv["status"] = "paid"
            inv["paid_at"] = datetime.now().isoformat()
            save_data(data)
            return True
    
    return False

def generate_invoice_docx(invoice_number):
    """Generate a DOCX invoice document."""
    data = load_data()
    
    # Find invoice
    invoice = None
    for inv in data["invoices"]:
        if inv["invoice_number"] == invoice_number:
            invoice = inv
            break
    
    if not invoice:
        raise ValueError(f"Invoice #{invoice_number} not found")
    
    company = data["company_info"]
    client = invoice["client"]
    
    # Create document
    doc = Document()
    
    # Header - Company Name
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(company["name"])
    run.bold = True
    run.font.size = Pt(18)
    
    # Company address
    addr = doc.add_paragraph()
    addr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    addr.add_run(f"{company['address']}\n{company['city_state_zip']}\n{company['phone']}\n{company['email']}")
    
    # Invoice title
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INVOICE")
    run.bold = True
    run.font.size = Pt(24)
    
    # Invoice details
    doc.add_paragraph()
    details = doc.add_table(rows=3, cols=2)
    details.cell(0, 0).text = "Invoice #:"
    details.cell(0, 1).text = str(invoice["invoice_number"])
    details.cell(1, 0).text = "Date:"
    details.cell(1, 1).text = datetime.fromisoformat(invoice["date"]).strftime("%B %d, %Y")
    details.cell(2, 0).text = "Due Date:"
    details.cell(2, 1).text = datetime.fromisoformat(invoice["due_date"]).strftime("%B %d, %Y")
    
    # Bill To
    doc.add_paragraph()
    bill_to = doc.add_paragraph()
    bill_to.add_run("BILL TO:\n").bold = True
    bill_to.add_run(f"{client['name']}\n{client['address']}\n{client['city_state_zip']}")
    
    # Line items table
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Rate'
    hdr_cells[3].text = 'Amount'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Line items
    for item in invoice["line_items"]:
        row = table.add_row().cells
        row[0].text = item["description"]
        row[1].text = str(item["quantity"])
        row[2].text = f"${item['rate']:,.2f}"
        row[3].text = f"${item['quantity'] * item['rate']:,.2f}"
    
    # Total row
    doc.add_paragraph()
    total_p = doc.add_paragraph()
    total_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = total_p.add_run(f"TOTAL: ${invoice['total']:,.2f}")
    run.bold = True
    run.font.size = Pt(14)
    
    # Notes
    if invoice.get("notes"):
        doc.add_paragraph()
        notes_p = doc.add_paragraph()
        notes_p.add_run("Notes:\n").bold = True
        notes_p.add_run(invoice["notes"])
    
    # Payment terms
    doc.add_paragraph()
    terms = doc.add_paragraph()
    terms.add_run("Payment Terms: ").bold = True
    terms.add_run("Net 30. Please make checks payable to " + company["name"])
    
    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filename = f"Invoice_{invoice_number}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    return filepath

def get_stats():
    """Get invoice statistics."""
    invoices = get_invoices()
    
    total_invoiced = sum(i.get("total", 0) for i in invoices)
    paid = [i for i in invoices if i.get("status") == "paid"]
    unpaid = [i for i in invoices if i.get("status") == "unpaid"]
    
    return {
        "total_invoices": len(invoices),
        "total_invoiced": total_invoiced,
        "paid_count": len(paid),
        "paid_amount": sum(i.get("total", 0) for i in paid),
        "unpaid_count": len(unpaid),
        "unpaid_amount": sum(i.get("total", 0) for i in unpaid)
    }
